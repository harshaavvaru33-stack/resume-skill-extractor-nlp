import re
import spacy
from pypdf import PdfReader
import docx

# Load spaCy model
nlp = spacy.load("en_core_web_sm")

# Comprehensive lexicon of skills categorized for dynamic matching
SKILLS_DB = {
    "Languages": [
        "python", "javascript", "typescript", "java", "c++", "c#", "go", 
        "rust", "php", "ruby", "swift", "kotlin", "scala", "r", "matlab", 
        "html", "css", "sql", "bash", "shell", "dart", "perl"
    ],
    "Frameworks & Libraries": [
        "react", "angular", "vue", "next.js", "nuxt.js", "svelte", 
        "django", "flask", "fastapi", "spring boot", "express", "nest.js", 
        "node.js", "laravel", "rails", "tensorflow", "pytorch", "keras", 
        "scikit-learn", "pandas", "numpy", "opencv", "bootstrap", "tailwind",
        "jquery", "redux", "graphql", "spring", "dotnet", "flask-restful"
    ],
    "Databases & Cache": [
        "postgresql", "mysql", "mongodb", "redis", "sqlite", "oracle", 
        "sql server", "dynamodb", "cassandra", "neo4j", "mariadb", 
        "elasticsearch", "firebase"
    ],
    "Cloud & DevOps": [
        "aws", "azure", "gcp", "docker", "kubernetes", "jenkins", 
        "terraform", "ansible", "github actions", "circleci", "heroku", 
        "netlify", "vercel", "vagrant", "prometheus", "grafana"
    ],
    "Tools & Methodologies": [
        "git", "github", "gitlab", "jira", "confluence", "figma", 
        "postman", "vscode", "linux", "unix", "nginx", "apache", 
        "agile", "scrum", "ci/cd", "rest api", "graphql", "microservices"
    ]
}

# Resume section headers for segmentation
SECTION_HEADERS = {
    "experience": ["experience", "work history", "employment", "professional background", "work experience", "career history", "professional experience"],
    "education": ["education", "academic profile", "academics", "educational background", "qualification", "qualifications", "academic background"],
    "projects": ["projects", "personal projects", "academic projects", "key projects", "featured projects"],
    "skills": ["skills", "technical skills", "core competencies", "skills & expertise", "technologies", "key skills"],
    "summary": ["summary", "professional summary", "about me", "objective", "profile"]
}

def extract_text_from_pdf(file_path):
    """Extract text from a PDF file."""
    text = ""
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
    return text

def extract_text_from_docx(file_path):
    """Extract text from a Word (docx) file."""
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
    return text

def segment_resume(text):
    """Segment resume text into distinct sections based on headers."""
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    sections = {}
    current_section = "contact_info"  # Start in contact info section
    sections[current_section] = []
    
    # Map headers to normalized section names
    header_mapping = {}
    for sec_name, headers in SECTION_HEADERS.items():
        for header in headers:
            header_mapping[header.lower()] = sec_name
            
    for line in lines:
        line_clean = line.strip().rstrip(':').strip()
        line_lower = line_clean.lower()
        
        is_header = False
        # Matches if the line is exactly a section header or a short line starting with a header
        if len(line_clean) < 40:
            for header, sec_name in header_mapping.items():
                if line_lower == header or line_lower.startswith(header + " "):
                    current_section = sec_name
                    if current_section not in sections:
                        sections[current_section] = []
                    is_header = True
                    break
        
        if not is_header:
            sections[current_section].append(line)
            
    return {k: "\n".join(v) for k, v in sections.items()}

def extract_name(text):
    """Extract the candidate's name using first-line heuristics and NER fallbacks."""
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    if not lines:
        return "Candidate Name"
        
    # Heuristic: Check first 3 lines for a candidate name
    for line in lines[:3]:
        clean_line = re.sub(r'\s+', ' ', line).strip()
        
        # Match capitalized word groups at the start (e.g. "Akash Kumar")
        match = re.match(r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)', clean_line)
        if match:
            name_candidate = match.group(1).strip()
            if not any(k in name_candidate.lower() for k in ["curriculum", "vitae", "resume", "summary"]):
                return name_candidate
                
        # Exclude contact details, URLs, or dividers
        if "@" in clean_line or any(char.isdigit() for char in clean_line):
            continue
        if any(sym in clean_line for sym in ["|", "/", "\\", "@", ":"]):
            continue
            
        words = clean_line.split()
        if 1 <= len(words) <= 4:
            lower_line = clean_line.lower()
            # Exclude section headers
            is_header = False
            for headers in SECTION_HEADERS.values():
                if lower_line in headers:
                    is_header = True
                    break
            if is_header:
                continue
                
            # Exclude technical terms
            is_skill = False
            for cat_skills in SKILLS_DB.values():
                if lower_line in [s.lower() for s in cat_skills]:
                    is_skill = True
                    break
            if is_skill:
                continue
                
            return clean_line
            
    # Fallback to NER if heuristics fail
    doc_first = nlp(text[:400])
    for ent in doc_first.ents:
        if ent.label_ == "PERSON":
            name = ent.text.strip()
            name = re.sub(r'\s+', ' ', name)
            words = name.split()
            if 1 <= len(words) <= 4:
                name_lower = name.lower()
                is_skill_or_noise = False
                for cat_skills in SKILLS_DB.values():
                    if name_lower in [s.lower() for s in cat_skills]:
                        is_skill_or_noise = True
                        break
                if not is_skill_or_noise and not any(k in name_lower for k in ["resume", "curriculum", "vitae", "email", "phone", "bangalore"]):
                    return name
                    
    return "Candidate Name"

def extract_email(text):
    """Extract email address using regex."""
    email_pattern = r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'
    emails = re.findall(email_pattern, text)
    return emails[0] if emails else ""

def extract_phone(text):
    """Extract phone number using regex."""
    # Matches patterns like +91 9999999999, (123) 456-7890, 123-456-7890, 9999999999
    phone_pattern = r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}|\+?\d{10,12}'
    phones = re.findall(phone_pattern, text)
    if phones:
        phone = phones[0].strip()
        return phone
    return ""

def extract_location(text):
    """Extract location (City, State, Country) using spaCy GPE entities and fallbacks."""
    doc = nlp(text[:800])
    gpe_entities = []
    
    # 1. Search for GPE/LOC entities in spaCy
    for ent in doc.ents:
        if ent.label_ in ["GPE", "LOC"]:
            val = ent.text.strip()
            val_lower = val.lower()
            # Ensure it isn't a tech skill (like Python or Java)
            is_skill = False
            for cat_skills in SKILLS_DB.values():
                if val_lower in [s.lower() for s in cat_skills]:
                    is_skill = True
                    break
            if not is_skill and len(val) > 2:
                gpe_entities.append(val)
                
    if gpe_entities:
        unique_gpes = []
        for gpe in gpe_entities:
            if gpe not in unique_gpes:
                unique_gpes.append(gpe)
        return ", ".join(unique_gpes[:2])
        
    # 2. Look for explicit city names or region patterns in first 8 lines
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    for line in lines[:8]:
        line_lower = line.lower()
        if any(kw in line_lower for kw in ["india", "usa", "uk", "canada", "california", "texas", "ny", "ca", "state", "city"]):
            match = re.search(r'\b([A-Z][a-zA-Z\s]+)\s*,\s*([A-Z][a-zA-Z\s]{1,15})\b', line)
            if match:
                return match.group(0)
                
        # Check common Indian/global cities as a fallback
        for city in ["bangalore", "bengaluru", "mumbai", "delhi", "pune", "hyderabad", "chennai", "kolkata", "noida", "gurgaon"]:
            if city in line_lower:
                match = re.search(rf'\b({city.title()}(?:\s*,\s*[A-Za-z]+)?)\b', line, re.IGNORECASE)
                if match:
                    return match.group(0)
                    
    return ""

def extract_links(text):
    """Extract LinkedIn, GitHub, and Portfolio URLs."""
    linkedin = re.findall(r'(https?://(?:www\.)?linkedin\.com/in/[a-zA-Z0-9_-]+)', text, re.IGNORECASE)
    github = re.findall(r'(https?://(?:www\.)?github\.com/[a-zA-Z0-9_-]+)', text, re.IGNORECASE)
    
    clean_lnk = lambda l: l.rstrip('.').rstrip('/')
    
    return {
        "linkedin": clean_lnk(linkedin[0]) if linkedin else "",
        "github": clean_lnk(github[0]) if github else ""
    }

def match_skills(text):
    """Case-insensitive skills matching with boundary handling for special chars."""
    text_lower = text.lower()
    matched = {}
    for category, skills in SKILLS_DB.items():
        matched[category] = []
        for skill in skills:
            escaped = re.escape(skill)
            if skill in ["c++", "c#"]:
                pattern = rf"(?<![a-zA-Z]){escaped}(?![a-zA-Z])"
            else:
                pattern = rf"\b{escaped}\b"
            
            if re.search(pattern, text_lower):
                matched[category].append(skill.title() if not skill.isupper() else skill)
    return matched

def parse_education(education_text):
    """Parse education credentials, degrees, institutions, and years."""
    education_list = []
    if not education_text.strip():
        return education_list
        
    lines = [line.strip() for line in education_text.split('\n') if line.strip()]
    
    degree_patterns = [
        r'\bB\.?Tech\b', r'\bM\.?Tech\b', r'\bB\.?E\.?\b', r'\bM\.?E\.?\b',
        r'\bB\.?Sc\b', r'\bM\.?Sc\b', r'\bPh\.?D\b', r'\bB\.?A\b', r'\bM\.?A\b',
        r'\bB\.?B\.?A\b', r'\bM\.?B\.?A\b', r'\bBachelor(?:s)?\b', r'\bMaster(?:s)?\b',
        r'\bDiploma\b', r'\bHSC\b', r'\bSSC\b'
    ]
    
    year_pattern = r'\b(19|20)\d{2}\b'
    
    for line in lines:
        degree_match = None
        for pat in degree_patterns:
            match = re.search(pat, line, re.IGNORECASE)
            if match:
                degree_match = match.group(0)
                break
        
        years = re.findall(year_pattern, line)
        year_str = " - ".join(years) if len(years) >= 2 else (years[0] if years else "")
        
        institution = ""
        inst_match = re.search(r'([A-Za-z0-9\s&,.-]+ (?:University|College|Institute|School|Academy|IIT|NIT|IIIT|BITS|Polytechnic)[A-Za-z0-9\s&,.-]*)', line, re.IGNORECASE)
        if inst_match:
            institution = inst_match.group(1).strip()
        else:
            doc = nlp(line)
            orgs = [ent.text for ent in doc.ents if ent.label_ == "ORG"]
            if orgs:
                institution = orgs[0]
                
        if degree_match or institution:
            # Format degree neatly
            degree_formatted = degree_match
            if degree_match:
                # Add subject if mentioned in line, e.g. "B.Tech in Computer Science"
                sub_match = re.search(rf'{re.escape(degree_match)}\s+(?:in|of)\s+([A-Za-z\s&]+)', line, re.IGNORECASE)
                if sub_match:
                    degree_formatted = f"{degree_match} in {sub_match.group(1).strip()}"
            
            education_list.append({
                "degree": degree_formatted if degree_formatted else "Degree / Qualification",
                "institution": institution if institution else "Institution Name",
                "duration": year_str if year_str else "Duration",
                "description": line if not degree_match and not institution else ""
            })
            
    if not education_list:
        # Fallback to paragraph block
        education_list.append({
            "degree": "Degree / Qualification",
            "institution": "University / School",
            "duration": "",
            "description": education_text.strip()
        })
        
    return education_list

def parse_experience(experience_text):
    """Parse work experiences including Job Title, Company, Dates, and Bullet Points."""
    experience_list = []
    if not experience_text.strip():
        return experience_list
        
    # Split by double newlines or lines containing clear dividers
    blocks = re.split(r'\n\n+', experience_text.strip())
    if len(blocks) <= 1:
        # Fallback: split by lines
        blocks = [line.strip() for line in experience_text.split('\n') if line.strip()]
        
    job_titles = [
        "Engineer", "Developer", "Manager", "Analyst", "Designer", "Consultant", "Intern", "Lead",
        "Architect", "Administrator", "Specialist", "Director", "Executive", "Scientist", "Programmer"
    ]
    
    date_regex = r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December|\d{1,2}/\d{2,4}|\d{4})\s*[-–to\s]+\s*(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December|\d{1,2}/\d{2,4}|\d{4}|Present))'
    
    for block in blocks:
        lines = [line.strip() for line in block.split('\n') if line.strip()]
        if not lines:
            continue
            
        header_line = lines[0]
        
        # Detect Job Title
        title = ""
        for jt in job_titles:
            if jt.lower() in header_line.lower():
                match = re.search(r'([A-Za-z0-9\s-]*' + jt + r'[A-Za-z0-9\s-]*)', header_line, re.IGNORECASE)
                if match:
                    title = match.group(1).strip()
                    break
        
        # Detect Date Range
        date_match = re.search(date_regex, header_line, re.IGNORECASE)
        date_str = date_match.group(1).strip() if date_match else ""
        if not date_str:
            years = re.findall(r'\b(?:19|20)\d{2}\b', header_line)
            if len(years) >= 2:
                date_str = f"{years[0]} - {years[1]}"
            elif len(years) == 1:
                date_str = f"{years[0]} - Present"
                
        # Detect Company
        company = ""
        doc = nlp(header_line)
        orgs = [ent.text for ent in doc.ents if ent.label_ == "ORG"]
        if orgs:
            company = orgs[0]
        else:
            at_match = re.search(r'\bat\s+([A-Z][A-Za-z0-9\s&]+)', header_line)
            if at_match:
                company = at_match.group(1).strip()
            else:
                parts = re.split(r'[-–|•,]', header_line)
                if len(parts) > 1:
                    potential_company = parts[1].strip()
                    if not any(d in potential_company.lower() for d in ["present", "201", "202"]):
                        company = potential_company
                        
        description_lines = lines[1:]
        description = "\n".join([re.sub(r'^[•\-\*\s]+', '', dl).strip() for dl in description_lines])
        
        if title or company or description:
            experience_list.append({
                "title": title if title else "Job Title",
                "company": company if company else "Company Name",
                "duration": date_str if date_str else "Duration",
                "description": description if description else ""
            })
            
    if not experience_list:
        experience_list.append({
            "title": "Professional Experience",
            "company": "Company Name",
            "duration": "",
            "description": experience_text.strip()
        })
        
    return experience_list

def parse_projects(projects_text):
    """Parse projects including title and description."""
    projects_list = []
    if not projects_text.strip():
        return projects_list
        
    blocks = re.split(r'\n\n+', projects_text.strip())
    if len(blocks) <= 1:
        blocks = [line.strip() for line in projects_text.split('\n') if line.strip()]
        
    for block in blocks:
        lines = [line.strip() for line in block.split('\n') if line.strip()]
        if not lines:
            continue
            
        title_line = lines[0]
        title = re.sub(r'^[•\-\*\s]+', '', title_line).strip()
        
        description = "\n".join([re.sub(r'^[•\-\*\s]+', '', l).strip() for l in lines[1:]])
        
        if title:
            projects_list.append({
                "title": title,
                "description": description
            })
            
    return projects_list

def parse_resume_text(text):
    """Main parsing orchestrator that converts raw text into structured JSON."""
    segmented = segment_resume(text)
    
    name = extract_name(text)
    email = extract_email(text)
    phone = extract_phone(text)
    location = extract_location(text)
    links = extract_links(text)
    skills = match_skills(text)
    
    # Parse sections
    education = parse_education(segmented.get("education", ""))
    experience = parse_experience(segmented.get("experience", ""))
    projects = parse_projects(segmented.get("projects", ""))
    
    # Extract summary/objective
    summary = segmented.get("summary", "").strip()
    if not summary:
        # Fallback: extract first paragraph if there's no experience or education in it
        lines = text.strip().split('\n')
        paragraph = []
        for line in lines[3:10]:  # Look right below name/contact
            if any(h in line.lower() for headers in SECTION_HEADERS.values() for h in headers):
                break
            if line.strip():
                paragraph.append(line.strip())
        if paragraph:
            summary = " ".join(paragraph)

    return {
        "personal_info": {
            "name": name,
            "email": email,
            "phone": phone,
            "location": location,
            "linkedin": links["linkedin"],
            "github": links["github"],
            "summary": summary
        },
        "skills": skills,
        "education": education,
        "experience": experience,
        "projects": projects
    }
